from docx import Document



def create_docx(text):


    filename = "Cover_Letter.docx"


    doc = Document()


    doc.add_heading(
        "Cover Letter",
        level=1
    )


    for line in text.split("\n"):

        doc.add_paragraph(
            line
        )


    doc.save(filename)


    return filename